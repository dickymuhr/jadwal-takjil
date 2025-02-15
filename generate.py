from docx import Document

from docx.shared import Pt
import re
import subprocess
from PyPDF2 import PdfReader

def count_pdf_pages(pdf_path):
    reader = PdfReader(pdf_path)
    return len(reader.pages)

def add_name_to_first_page_header(document, search_name):
    tables = document.tables
    if tables:
        first_table = tables[0]  # Get the first table
        parent = first_table._element.getparent()  # Get the table's parent XML element
        
        # Find the first empty paragraph after the table
        for index, para in enumerate(document.paragraphs):
            if para.text.strip() == "":  # Check for an empty paragraph
                new_paragraph = document.paragraphs[index]  # Use the same paragraph to replace text
                new_paragraph.text = search_name  # Replace empty line with name
                
                # Apply formatting
                new_paragraph.alignment = 2  # Right-aligned text
                run = new_paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True

                # print("✅ Successfully replaced empty line with name.")
                return  # Stop after replacing the first empty line

    print("⚠️ No empty paragraph found after the first table.")


def bold_name_in_tables(document_path, search_name, output_path):
    doc = Document(document_path)
    pattern = re.compile(re.escape(search_name), re.IGNORECASE)

    add_name_to_first_page_header(doc, search_name)

    found = False
    count_bolded = 0

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if pattern.search(cell.text):
                    found = True
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        paragraph.clear()
                        start = 0
                        for match in pattern.finditer(original_text):
                            if match.start() > start:
                                non_bold_run = paragraph.add_run(original_text[start:match.start()])
                                non_bold_run.font.name = 'Times New Roman'  # Set font for non-bold text
                            bold_run = paragraph.add_run(original_text[match.start():match.end()])
                            bold_run.bold = True
                            bold_run.font.name = 'Times New Roman'
                            bold_run.font.size = Pt(12)
                            count_bolded += 1
                            start = match.end()
                        if start < len(original_text):
                            remaining_text_run = paragraph.add_run(original_text[start:])
                            remaining_text_run.font.name = 'Times New Roman'
    
    if found:
        doc.save(output_path)
        print(f"Name found and bolded {count_bolded} times.", search_name)
        return True, count_bolded
    else:
        print("Name not found.", search_name)
        doc.save(output_path)
        return False, 0

def convert_docx_to_pdf_libreoffice(input_path, output_path):
    try:
        cmd = ['libreoffice', '--convert-to', 'pdf', '--outdir', output_path, input_path]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        print("Conversion successful.")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")

# # Example usage
# document_path = 'takjil jabur fix 1445.docx'  # Ensure the path is correct
# search_name = 'Ibu Hartiono Wahyudi'
# output_docx_path = f'/home/diccode/jadwal-takjil/docx/{search_name}.docx'
# found, count_bolded = bold_name_in_tables(document_path, search_name, output_docx_path)

# if found:
#     output_pdf_path = '/home/diccode/jadwal-takjil/pdf'  # Ensure the output path is correct
#     convert_docx_to_pdf_libreoffice(output_docx_path, output_pdf_path)

def process_names_from_file(name_file_path, document_path, docx_output_base_path, pdf_output_base_path, expected_page_count):
    with open(name_file_path, 'r') as file:
        names = file.readlines()
    
    # Iterate over each name and process it
    for name in names:
        name = name.strip()  # Remove any leading/trailing whitespace
        if name:  # Check if name is not empty
            # Define output paths for the DOCX and PDF files for this name
            output_docx_path = f"{docx_output_base_path}/{name}.docx"
            output_pdf_path = f"{pdf_output_base_path}"
            
            # Bold the name in the tables and save the document
            found, count_bolded = bold_name_in_tables(document_path, name, output_docx_path)
            
            # if found:
            #     # Convert the modified DOCX to PDF
            #     convert_docx_to_pdf_libreoffice(output_docx_path, output_pdf_path)
                
            #     # Perform page count check here after conversion
            #     actual_page_count = count_pdf_pages(output_docx_path.replace('docx','pdf'))
            #     if actual_page_count != expected_page_count:
            #         error_message = f"Error: The document for '{name}' was expected to have {expected_page_count} pages, but it has {actual_page_count}."
            #         raise ValueError(error_message)
                
            #     print(f"Document for '{name}' is correctly formatted with {expected_page_count} pages.")
            # else:
            #     error_message = f"No occurrences found for {name}."
            #     raise ValueError(error_message)

# Example usage
name_file_path = '/home/diccode/jadwal-takjil/name_list.txt'  # Path to your .txt file containing names
document_path = '/home/diccode/jadwal-takjil/FIX JADWAL PEMBERI HIDANGAN 1446 H.docx'  # Path to the source document
docx_output_base_path = '/home/diccode/jadwal-takjil/docx'  # Base path where modified DOCX files will be saved
pdf_output_base_path = '/home/diccode/jadwal-takjil/pdf'  # Base path where converted PDF files will be saved

expected_page_count = 7 
process_names_from_file(name_file_path, document_path, docx_output_base_path, pdf_output_base_path, expected_page_count)