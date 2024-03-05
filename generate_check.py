from docx import Document
from docx.shared import Pt
import re
import subprocess

def delete_name_in_tables(document_path, search_name, output_path):
    doc = Document(document_path)
    pattern = re.compile(re.escape(search_name), re.IGNORECASE)
    found = False
    count_deleted = 0

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = pattern.finditer(cell.text)
                if any(matches):  # Check if there are any matches in the cell
                    found = True
                    for paragraph in cell.paragraphs:
                        new_paragraph_text = pattern.sub('', paragraph.text)  # Delete matched text
                        count_deleted += len(list(pattern.finditer(paragraph.text)))  # Count occurrences
                        paragraph.clear()
                        paragraph.add_run(new_paragraph_text)  # Add updated text without the deleted name
                    
    if found:
        doc.save(output_path)
        print(f"Name '{search_name}' found and deleted {count_deleted} times.")
        return True, count_deleted
    else:
        print(f"Name '{search_name}' not found.")
        return False, 0

# Adjust the example usage accordingly
document_path = 'takjil jabur fix 1445 copy.docx'  # Ensure the path is correct
name_file_path = 'name_list.txt'  # Path to your .txt file containing names

# Iterate over each name and delete it from the document
def process_names_and_delete_from_file(name_file_path, document_path):
    with open(name_file_path, 'r') as file:
        names = [name.strip() for name in file.readlines() if name.strip()]

    for name in names:
        # Use the same document_path as both input and output to modify the document in place
        found, count_deleted = delete_name_in_tables(document_path, name, document_path)
        if not found:
            print(f"No occurrences found for {name}.")

# Call the function with the updated paths
process_names_and_delete_from_file(name_file_path, document_path)
